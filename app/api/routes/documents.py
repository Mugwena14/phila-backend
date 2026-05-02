from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from uuid import UUID
from typing import List, Optional
from pydantic import BaseModel
import json
import uuid
import os
import re
import shutil

from app.db.database import get_db
from app.models.patient_document import PatientDocument
from app.models.document_template import DocumentTemplate
from app.models.user import User
from app.models.doctor import Doctor
from app.models.booking import Booking
from app.core.security import decode_token
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import logging


from fastapi.responses import Response as FastAPIResponse
from app.services.template_starters import create_starter_template

router = APIRouter(prefix="/documents", tags=["documents"])
security = HTTPBearer()
logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads/templates"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: Session = Depends(get_db),
) -> User:
    token = credentials.credentials
    payload = decode_token(token)
    if payload is None:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    user = db.query(User).filter(User.id == payload.get("sub")).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user


def extract_placeholders(file_path: str) -> list[str]:
    """Extract all {{placeholder}} keys from a .docx file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    placeholders = re.findall(r'\{\{(\w+)\}\}', text)
    return list(dict.fromkeys(placeholders))  # deduplicate, preserve order


def fill_placeholders(template_path: str, values: dict, output_path: str) -> None:
    """Replace {{placeholder}} with actual values in a .docx and save."""
    from docx import Document as DocxDocument

    doc = DocxDocument(template_path)

    def replace_in_paragraph(para):
        for run in para.runs:
            for key, val in values.items():
                run.text = run.text.replace(f"{{{{{key}}}}}", str(val))

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(output_path)


# ── TEMPLATE ENDPOINTS ────────────────────────────────────────────

@router.post("/templates/upload", status_code=201)
async def upload_template(
    name: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a .docx template file. Extracts {{placeholders}} automatically."""
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    if not doctor:
        raise HTTPException(status_code=404, detail="Doctor profile not found")

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    # Save file
    template_id = str(uuid.uuid4())
    safe_filename = f"{template_id}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Extract placeholders
    try:
        placeholders = extract_placeholders(file_path)
    except Exception as e:
        os.remove(file_path)
        raise HTTPException(status_code=400, detail=f"Could not read .docx file: {str(e)}")

    if not placeholders:
        logger.warning(f"No placeholders found in template {file.filename}")

    # Save to DB
    template = DocumentTemplate(
        id=uuid.UUID(template_id),
        doctor_id=doctor.id,
        name=name,
        description=description,
        filename=file.filename,
        file_path=file_path,
        placeholders=json.dumps(placeholders),
    )
    db.add(template)
    db.commit()
    db.refresh(template)

    logger.info(f"Template uploaded: {name} — {len(placeholders)} placeholders found")

    return {
        "id": str(template.id),
        "name": template.name,
        "description": template.description,
        "filename": template.filename,
        "placeholders": placeholders,
        "created_at": str(template.created_at),
    }


@router.get("/templates")
def list_templates(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all templates for this doctor."""
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    if not doctor:
        raise HTTPException(status_code=404, detail="Doctor profile not found")

    templates = db.query(DocumentTemplate).filter(
        DocumentTemplate.doctor_id == doctor.id
    ).order_by(DocumentTemplate.created_at.desc()).all()

    return [
        {
            "id": str(t.id),
            "name": t.name,
            "description": t.description,
            "filename": t.filename,
            "placeholders": json.loads(t.placeholders),
            "created_at": str(t.created_at),
        }
        for t in templates
    ]


@router.delete("/templates/{template_id}", status_code=200)
def delete_template(
    template_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    if not doctor:
        raise HTTPException(status_code=404, detail="Doctor profile not found")

    template = db.query(DocumentTemplate).filter(
        DocumentTemplate.id == template_id,
        DocumentTemplate.doctor_id == doctor.id,
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # Delete file
    if os.path.exists(template.file_path):
        os.remove(template.file_path)

    db.delete(template)
    db.commit()
    return {"message": "Template deleted"}


@router.post("/templates/{template_id}/generate")
def generate_from_template(
    template_id: UUID,
    booking_id: UUID,
    values: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Fill template placeholders with provided values.
    Returns filled .docx as download.
    """
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    if not doctor:
        raise HTTPException(status_code=404, detail="Doctor profile not found")

    template = db.query(DocumentTemplate).filter(
        DocumentTemplate.id == template_id,
        DocumentTemplate.doctor_id == doctor.id,
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    booking = db.query(Booking).filter(Booking.id == booking_id).first()
    if not booking:
        raise HTTPException(status_code=404, detail="Booking not found")

    # Generate output file
    output_filename = f"filled_{uuid.uuid4()}.docx"
    output_path = os.path.join(UPLOAD_DIR, output_filename)

    try:
        fill_placeholders(template.file_path, values, output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error filling template: {str(e)}")

    # Save record to patient_documents
    doc = PatientDocument(
        id=uuid.uuid4(),
        patient_id=booking.patient_id,
        booking_id=booking.id,
        doctor_id=doctor.id,
        doc_type=f"template_{template_id}",
        content=json.dumps({
            "_template_id": str(template_id),
            "_template_name": template.name,
            "_output_file": output_path,
            **values,
        }),
        generated_by=current_user.id,
    )
    db.add(doc)
    db.commit()

    # Return the filled .docx as a download
    return FileResponse(
        path=output_path,
        filename=f"{template.name}_{booking_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ── EXISTING ENDPOINTS (unchanged) ───────────────────────────────

class GenerateDocumentRequest(BaseModel):
    booking_id: UUID
    doc_type: str
    content: dict


@router.post("/generate", status_code=201)
def generate_document(
    data: GenerateDocumentRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    if not doctor:
        raise HTTPException(status_code=404, detail="Doctor profile not found")

    booking = db.query(Booking).filter(Booking.id == data.booking_id).first()
    if not booking:
        raise HTTPException(status_code=404, detail="Booking not found")

    doc = PatientDocument(
        id=uuid.uuid4(),
        patient_id=booking.patient_id,
        booking_id=booking.id,
        doctor_id=doctor.id,
        doc_type=data.doc_type,
        content=json.dumps(data.content),
        generated_by=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "id": str(doc.id),
        "doc_type": doc.doc_type,
        "content": json.loads(doc.content),
        "created_at": str(doc.created_at),
        "message": "Document generated successfully",
    }


@router.get("/patient/{patient_id}")
def get_patient_documents(
    patient_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    docs = db.query(PatientDocument).filter(
        PatientDocument.patient_id == patient_id
    ).order_by(PatientDocument.created_at.desc()).all()

    return [
        {
            "id": str(d.id),
            "patient_id": str(d.patient_id),
            "booking_id": str(d.booking_id) if d.booking_id else None,
            "doc_type": d.doc_type,
            "content": json.loads(d.content),
            "created_at": str(d.created_at),
        }
        for d in docs
    ]


@router.get("/{doc_id}")
def get_document(
    doc_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(PatientDocument).filter(PatientDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "id": str(doc.id),
        "doc_type": doc.doc_type,
        "content": json.loads(doc.content),
        "created_at": str(doc.created_at),
    }


@router.get("/templates/starter/{doc_type}")
def download_starter_template(
    doc_type: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download a pre-built starter .docx template for a given doc type."""
    valid_types = ['sick_letter', 'medical_certificate', 'referral_letter', 'visit_summary']
    if doc_type not in valid_types:
        raise HTTPException(status_code=400, detail="Invalid doc type")

    # Get doctor info to pre-fill the letterhead
    doctor = db.query(Doctor).filter(Doctor.user_id == current_user.id).first()
    practice_name = doctor.practice_name if doctor else "Your Practice Name"

    user = db.query(User).filter(User.id == current_user.id).first()
    doctor_name = f"Dr. {user.full_name}" if user else "Dr. Your Name"

    docx_bytes = create_starter_template(doc_type, practice_name, doctor_name)

    filename_map = {
        'sick_letter': 'Sick_Letter_Starter.docx',
        'medical_certificate': 'Medical_Certificate_Starter.docx',
        'referral_letter': 'Referral_Letter_Starter.docx',
        'visit_summary': 'Visit_Summary_Starter.docx',
    }

    return FastAPIResponse(
        content=docx_bytes,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename_map[doc_type]}"'}
    )