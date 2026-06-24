"""
Document template service - extraction and substitution for .docx templates.

Handles the messy reality of real-world Word documents:
  - Smart quotes from Word autocorrect that silently break placeholder regex
  - Placeholders split across multiple <w:r> runs (the python-docx classic gotcha)
  - Placeholders in headers and footers, not just the body
  - Validation feedback for empty or unreadable files
"""
from typing import List, Dict, Tuple
import re
import logging

logger = logging.getLogger(__name__)

# Matches {{placeholder_name}} - the standard syntax we ask doctors to use.
# \w+ allows letters, digits, underscores. No spaces inside placeholders.
PLACEHOLDER_PATTERN = re.compile(r"\{\{(\w+)\}\}")


def _normalise_smart_chars(text: str) -> str:
    """
    Word autocorrect helpfully replaces straight quotes and braces with curly
    Unicode variants. The placeholder regex needs ASCII { and }, so any
    Unicode variant of { } " ' must be flipped back to ASCII before matching.

    Idempotent: running this on already-normalised text yields the same output.
    """
    if not text:
        return text
    replacements = {
        # Curly single quotes -> straight
        "\u2018": "'",
        "\u2019": "'",
        "\u201A": "'",
        "\u201B": "'",
        # Curly double quotes -> straight
        "\u201C": '"',
        "\u201D": '"',
        "\u201E": '"',
        "\u201F": '"',
        # Fullwidth brace variants -> ASCII
        "\uFF5B": "{",
        "\uFF5D": "}",
        # White-square braces (rare but seen)
        "\u2774": "{",
        "\u2775": "}",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def _normalise_paragraph_runs(para) -> None:
    """
    Apply smart-char normalisation in-place to every run in a paragraph.
    Done at the run level so we don't lose formatting.
    """
    for run in para.runs:
        if run.text:
            run.text = _normalise_smart_chars(run.text)


def _iter_all_paragraphs(doc):
    """
    Yield every paragraph in the document - body, tables, and all section
    headers and footers. python-docx treats headers/footers as separate
    section objects that are easy to miss.
    """
    # Body paragraphs
    for para in doc.paragraphs:
        yield para
    # Table cells (which contain their own paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    # Section headers and footers
    for section in doc.sections:
        if section.header is not None:
            for para in section.header.paragraphs:
                yield para
            # Tables inside headers
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para
        if section.footer is not None:
            for para in section.footer.paragraphs:
                yield para
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para


def extract_placeholders(file_path: str) -> List[str]:
    """
    Scan a .docx and return the ordered list of unique {{placeholder}} names.
    Covers body, tables, headers, and footers.
    """
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        logger.warning(f"Could not open .docx for extraction at {file_path}: {e}")
        raise ValueError(f"Could not read .docx file: {e}")

    seen: Dict[str, None] = {}  # ordered set
    for para in _iter_all_paragraphs(doc):
        # Use the full paragraph text (joined across runs) so split placeholders
        # are still detected. We do not modify the document here.
        text = _normalise_smart_chars(para.text)
        for match in PLACEHOLDER_PATTERN.findall(text):
            seen[match] = None
    return list(seen.keys())


def _replace_in_paragraph(para, values: Dict[str, str]) -> None:
    """
    Replace {{placeholders}} in a single paragraph.

    The hard part: python-docx splits text across runs. A placeholder typed
    as {{patient_name}} and then mid-edited can end up with `{{patient` in one
    run and `_name}}` in the next, so a per-run str.replace silently misses it.

    Approach: if any placeholder exists in the joined paragraph text but the
    run-by-run replace alone wouldn't catch it, collapse all runs into the
    first run and clear the others. This loses run-level formatting WITHIN the
    placeholder substring, but preserves formatting of the rest of the paragraph
    in practice because Word usually wraps the whole placeholder in one
    formatting span.
    """
    if not para.runs:
        return

    # First, normalise smart chars in every run so the regex finds them.
    _normalise_paragraph_runs(para)

    # Try the simple per-run replace first - covers the easy case and keeps
    # all formatting intact.
    for run in para.runs:
        if run.text and "{{" in run.text:
            for key, val in values.items():
                token = "{{" + key + "}}"
                if token in run.text:
                    run.text = run.text.replace(token, str(val))

    # Check if any placeholders are still present in the joined text.
    # If yes, they were split across runs. Collapse and re-replace.
    joined = "".join(run.text for run in para.runs)
    if "{{" not in joined:
        return

    still_present = [k for k in values.keys() if ("{{" + k + "}}") in joined]
    if not still_present:
        return

    # Substitute in the joined string, then put it all in the first run.
    new_text = joined
    for key in still_present:
        new_text = new_text.replace("{{" + key + "}}", str(values[key]))

    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def fill_template(template_path: str, values: Dict[str, str], output_path: str) -> None:
    """
    Open the template, substitute all {{placeholders}} with provided values,
    save to output_path. Covers body, tables, headers, footers.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(template_path)
    for para in _iter_all_paragraphs(doc):
        _replace_in_paragraph(para, values)
    doc.save(output_path)


def build_preview_values(placeholders: List[str]) -> Dict[str, str]:
    """
    Build a dummy values dict for previewing a template against sample data.
    Pattern-matches placeholder names to plausible sample content; falls back
    to "[placeholder_name]" so the doctor can clearly see which field is which
    in the rendered output.
    """
    today_str = "12 March 2026"
    values: Dict[str, str] = {}
    for key in placeholders:
        k = key.lower()
        if "patient" in k and "name" in k:
            values[key] = "Sipho Mthembu"
        elif "doctor" in k:
            values[key] = "Dr. Jane Smith"
        elif "practice" in k:
            values[key] = "Mthembu Family Practice"
        elif "date" in k and "visit" in k:
            values[key] = today_str
        elif "date" in k:
            values[key] = today_str
        elif "diagnosis" in k or "concern" in k:
            values[key] = "Upper respiratory tract infection"
        elif "medication" in k:
            values[key] = "Amoxicillin 500mg, Paracetamol 500mg"
        elif "allerg" in k:
            values[key] = "Penicillin"
        elif "note" in k or "additional" in k:
            values[key] = "Patient advised to rest and increase fluid intake."
        elif "duration" in k:
            values[key] = "3 days"
        elif "severity" in k:
            values[key] = "5/10"
        elif "days_off" in k or ("days" in k and "off" in k):
            values[key] = "3"
        elif "qualification" in k:
            values[key] = "MBChB (UCT)"
        elif "hpcsa" in k:
            values[key] = "PR1234567"
        elif "urgency" in k:
            values[key] = "Routine"
        elif "referred" in k and "specialty" in k:
            values[key] = "Cardiologist"
        elif "history" in k:
            values[key] = "Hypertension diagnosed 2024, well-controlled on Amlodipine."
        else:
            # Visible placeholder so the doctor can see which slot it is
            values[key] = f"[{key}]"
    return values
