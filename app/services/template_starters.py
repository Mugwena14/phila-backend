from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def create_starter_template(doc_type: str, practice_name: str = "Your Practice Name", doctor_name: str = "Dr. Your Name") -> bytes:
    """Generate a starter .docx template with placeholders pre-inserted."""
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Helper: add a styled paragraph ───────────────────────────
    def add_para(text='', bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_field_line(label: str, placeholder: str, size=11):
        """Add a line like:  Patient name:  {{patient_name}}"""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        label_run = p.add_run(f"{label}:  ")
        label_run.bold = True
        label_run.font.size = Pt(size)
        field_run = p.add_run(placeholder)
        field_run.font.size = Pt(size)
        field_run.font.color.rgb = RGBColor(15, 118, 110)  # teal — easy to spot
        return p

    # ── PRACTICE LETTERHEAD ───────────────────────────────────────
    add_para(practice_name, bold=True, size=16, color=(15, 118, 110), align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doctor_name, size=11, color=(80, 80, 80))
    add_para('{{qualification}}', size=10, color=(15, 118, 110))
    add_para('HPCSA: {{hpcsa_number}}', size=10, color=(80, 80, 80))
    add_para('{{practice_address}}', size=10, color=(80, 80, 80))

    # Divider
    doc.add_paragraph('─' * 60)

    if doc_type == 'sick_letter':
        add_para('SICK CERTIFICATE', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        add_field_line('Date issued', '{{date_issued}}')
        doc.add_paragraph()
        add_para('To whom it may concern,', size=11, space_after=10)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run('This is to certify that ').font.size = Pt(11)
        r = p.add_run('{{patient_name}}'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(15,118,110); r.bold = True
        p.add_run(' was examined at this practice on ').font.size = Pt(11)
        r2 = p.add_run('{{date_of_visit}}'); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(15,118,110)
        p.add_run(' and is unfit for work/school for ').font.size = Pt(11)
        r3 = p.add_run('{{days_off}}'); r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(15,118,110)
        p.add_run(' day(s), from ').font.size = Pt(11)
        r4 = p.add_run('{{from_date}}'); r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(15,118,110)
        p.add_run(' to ').font.size = Pt(11)
        r5 = p.add_run('{{to_date}}').font.size = Pt(11)

        doc.add_paragraph()
        add_field_line('Reason / Diagnosis', '{{diagnosis}}')
        add_field_line('Additional notes', '{{notes}}')
        doc.add_paragraph()
        add_para('Yours sincerely,', size=11, space_after=20)
        add_para('_______________________', size=11)
        add_field_line('Name', '{{doctor_name}}')
        add_field_line('Practice', '{{practice_name}}')

    elif doc_type == 'medical_certificate':
        add_para('MEDICAL CERTIFICATE', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        add_field_line('Date issued', '{{date_issued}}')
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run('I, ').font.size = Pt(11)
        r = p.add_run('{{doctor_name}}'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(15,118,110); r.bold = True
        p.add_run(' (').font.size = Pt(11)
        r2 = p.add_run('{{qualification}}'); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(15,118,110)
        p.add_run('), hereby certify that I examined ').font.size = Pt(11)
        r3 = p.add_run('{{patient_name}}'); r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(15,118,110); r3.bold = True
        p.add_run(' on ').font.size = Pt(11)
        r4 = p.add_run('{{date_of_visit}}'); r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(15,118,110)
        p.add_run('.').font.size = Pt(11)

        doc.add_paragraph()
        add_field_line('Diagnosis', '{{diagnosis}}')
        add_field_line('Duration of illness', '{{duration}}')
        add_field_line('Notes', '{{notes}}')
        doc.add_paragraph()
        add_para('_______________________', size=11, space_after=4)
        add_field_line('Signature', '{{doctor_name}}')
        add_field_line('HPCSA number', '{{hpcsa_number}}')
        add_field_line('Practice', '{{practice_name}}')

    elif doc_type == 'referral_letter':
        add_para('REFERRAL LETTER', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        add_field_line('Date', '{{date_issued}}')
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run('Dear Dr. ').font.size = Pt(11)
        r = p.add_run('{{referred_to_doctor}}'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(15,118,110)
        p.add_run(' / ').font.size = Pt(11)
        r2 = p.add_run('{{referred_to_specialty}}'); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(15,118,110)
        p.add_run(' Specialist,').font.size = Pt(11)

        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        p2.add_run('I am referring my patient, ').font.size = Pt(11)
        r3 = p2.add_run('{{patient_name}}'); r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(15,118,110); r3.bold = True
        p2.add_run(', for specialist assessment.').font.size = Pt(11)

        doc.add_paragraph()
        add_field_line('Reason for referral', '{{reason_for_referral}}')
        add_field_line('Relevant history', '{{relevant_history}}')
        add_field_line('Current medications', '{{current_medications}}')
        add_field_line('Allergies', '{{allergies}}')
        add_field_line('Urgency', '{{urgency}}')
        doc.add_paragraph()
        add_para('Kind regards,', size=11, space_after=20)
        add_para('_______________________', size=11, space_after=4)
        add_field_line('Referring doctor', '{{referring_doctor}}')
        add_field_line('Practice', '{{practice_name}}')

    elif doc_type == 'visit_summary':
        add_para('VISIT SUMMARY', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        add_field_line('Patient', '{{patient_name}}')
        add_field_line('Date of visit', '{{date_of_visit}}')
        add_field_line('Doctor', '{{doctor_name}}')
        add_field_line('Practice', '{{practice_name}}')
        doc.add_paragraph()
        add_para('CLINICAL NOTES', bold=True, size=11, color=(80,80,80))
        add_field_line('Chief complaint', '{{chief_complaint}}')
        add_field_line('Duration', '{{duration}}')
        add_field_line('Severity', '{{severity}}')
        add_field_line('Medications prescribed', '{{medications_prescribed}}')
        add_field_line('Allergies', '{{allergies}}')
        doc.add_paragraph()
        add_para('PLAN', bold=True, size=11, color=(80,80,80))
        add_field_line('Recommendations', '{{recommendations}}')
        add_field_line('Follow-up', '{{follow_up}}')
        add_field_line('Additional notes', '{{notes}}')

    # ── INSTRUCTION NOTE AT THE BOTTOM ───────────────────────────
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    nr = note.add_run('HOW TO USE THIS TEMPLATE: ')
    nr.bold = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(120, 120, 120)
    nr2 = note.add_run(
        'The teal {{placeholder}} fields will be filled automatically from patient intake data. '
        'Customize the layout, add your logo, change fonts and colours — but keep the {{placeholders}} exactly as they are. '
        'Delete this instruction block before uploading. Save as .docx and upload to Phila.'
    )
    nr2.font.size = Pt(9)
    nr2.font.color.rgb = RGBColor(150, 150, 150)

    # ── Return as bytes ───────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()